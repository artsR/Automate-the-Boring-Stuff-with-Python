#! python3
# InvitationForm.py - 

import docx

guestsFile = open('guests.txt', 'r')
guestNames = guestsFile.readlines()

doc = docx.Document('invitationForm.docx')

for name in guestNames:
    para1 = doc.add_paragraph('It would be a pleasure to have the company of')
    para1.style = 'Python_A'
    para2 = doc.add_paragraph(name.strip('\n'))
    para2.style = 'Python_B'
    para3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    para3.style = 'Python_A'
    para4 = doc.add_paragraph('April 1st')
    para4.style = 'Python_C'
    para5 = doc.add_paragraph('at 7 o\'clock')
    para5.style = 'Python_A'
    para5.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.save('Invitation.docx')
