########### CHAPTER 13 ###########

import docx # refers to 'python-docx' Module

'''
    Compared to plaintext, .docx files have a lot of structure. This structure
is represented by three different data types in Python-Docx.

    1. 'Document' object - represents the entire document
    2. 'Paragraph' object - paragraphs in the document.
                New paragraph begins whenever pressed ENTER or RETURN
                Each of these Paragraph objects contains a list of one or more:
    3. 'Run' objects - contiguous run of text with the same style.
                      https://gyazo.com/93604ccd387ceb66ce1ddb8f8fde340d

                                                                            '''
# https://python-docx.readthedocs.org/


filename = 'Chapter 13.docx'
doc = docx.Document(filename) # returns 'Document' object

print(len(doc.paragraphs)) # number of paragraph in the file (no. of elem. of list)
doc.paragraphs[0].text
print(doc.paragraphs[1].text) # all text in this paragraph w/o taking into 
    #account its style formatting
print(len(doc.paragraphs[1].runs)) # number of runs in paragraph[1]
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text

# To get all text from document, ignoring style I may write 'getText(filename)'
#This function should look like this:
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text) #or fullText.append('    ' + para.text)
    return '\n'.join(fullText)

# Run attributes:
'''
        Attribute:                           Description:
        
            bold              The text appears in bold.
            
            italic            The text appears in italic.
            
            underline         The text is underlined.
            
            strike            The text appears with strikethrough.
            
            double_strike     The text appears with double strikethrough.
            
            all_caps          The text appears in capital letters.
            
            small_caps        The text appears in capital letters, with
                                lowercase letters two points smaller.
                                
            shadow            The text appears with a shadow.
            
            outline           The text appears outlined rather than solid.
            
            rtl               The text is written right-to-left.
            
            imprint           The text appears pressed into the page.
            
            emboss            The text appears raised off the page in relief

https://python-docx.readthedocs.org/en/latest/user/styles.html
                                                                            '''
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
 doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('Chapter 13_reStyle.docx')

# Writing Word Document
doc = docx.Document() # returns a new, blank Word Document object.
doc.add_paragraph('Hello Word!', 'Title') # adds a new paragraph of text
        #to the document and returns a reference to the 'Paragraph' object
        #that was added. 'Title' Style used to this paragraph.
paraObj1 = doc.add_paragraph('This is a Second Paragraph.')
paraObj2 = doc.add_paragraph('Third Paragraph.')
paraObj1.add_run(' I\'m excited.') # it will be added into 'Second Paragraph'
doc.save('helloWord.docx')

# add_heading - adds a new paragraph with determined style:
doc.add_heading('Text with Header 0', 0) # 0 is a 'Title' Style
doc.add_heading('Text with Header 1', 1) # 1 is a Main Heading
doc.add_heading('Text with Header 4', 4) # 4 is a lowest Subheading
doc.save('helloWord.docx')

# Adding Line and Page Breaks:
paraObj1.runs[1].add_break() # adds line break after 'runs[1]'
paraObj1.add_run('Now, I am after the "Line Break".') #this is Run[2]
paraObj2.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
paraObj2.add_run('Now, I am on a New PAGE, after "Page Break".')#creates New Page
doc.save('helloWord.docx')

# Adding Pictures: (can be added only to the end of document)
doc.add_picture('Chapter 13.png', width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))
